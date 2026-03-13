#!/usr/bin/env python3
"""
EPUB Converter Server
Run: python app.py
Open: http://localhost:5000
Requires: pip install flask ebooklib python-docx markdown2
"""

import os, re, datetime, threading, uuid
from pathlib import Path
from flask import Flask, request, send_file, jsonify, render_template_string

# ── Optional deps check ───────────────────────────────────────────────────────
try:
    from ebooklib import epub
    from docx import Document as DocxDocument
    import markdown2
except ImportError as e:
    print(f"\n[ERROR] Missing package: {e}")
    print("Run:  pip install flask ebooklib python-docx markdown2\n")
    raise

app = Flask(__name__)
UPLOAD_FOLDER = Path("/tmp/epub_converter")
UPLOAD_FOLDER.mkdir(exist_ok=True)

# Auto-cleanup files older than 10 minutes
def _cleanup():
    now = datetime.datetime.now().timestamp()
    for f in UPLOAD_FOLDER.iterdir():
        if now - f.stat().st_mtime > 600:
            f.unlink(missing_ok=True)
    threading.Timer(120, _cleanup).start()
_cleanup()


# ── Parsers ───────────────────────────────────────────────────────────────────

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def detect_language(text: str) -> str:
    hebrew = len(re.findall(r'[\u0590-\u05FF]', text))
    latin  = len(re.findall(r'[A-Za-z]', text))
    if hebrew > 0 and latin == 0: return "he"
    if latin  > 0 and hebrew == 0: return "en"
    if hebrew > 0 and latin  > 0: return "mixed"
    return "en"

def xml_para_text(para) -> str:
    """
    Extract all visible text from a paragraph XML node, including text inside
    tracked-change insertions (w:ins), comment anchors, and revision markup —
    anything that holds a w:t element. Skips deleted text (w:del).
    """
    root = para._p
    parts = []
    for el in root.iter():
        # Skip deleted runs entirely
        if el.tag == f"{{{WNS}}}del":
            continue
        if el.tag == f"{{{WNS}}}t":
            t = (el.text or "")
            # Check parent isn't inside a w:del
            parts.append(t)
    return "".join(parts).strip()

def xml_para_style(para) -> str:
    try:
        return para.style.name if para.style and para.style.name else ""
    except Exception:
        return ""

def xml_outline_level(para) -> int:
    try:
        pPr = para._p.find(f"{{{WNS}}}pPr")
        if pPr is None: return 99
        ol = pPr.find(f"{{{WNS}}}outlineLvl")
        if ol is None: return 99
        return int(ol.get(f"{{{WNS}}}val", 99))
    except Exception:
        return 99

def xml_run_formatting(run) -> tuple:
    """Return (bold, italic) for a run."""
    try:
        return bool(run.bold), bool(run.italic)
    except Exception:
        return False, False

def para_to_html(para) -> str:
    """
    Build HTML for a paragraph. Tries run-level formatting first;
    falls back to raw XML text extraction.
    """
    # Try runs first (gives bold/italic)
    parts = []
    for run in para.runs:
        t = (run.text or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        if not t: continue
        b, i = xml_run_formatting(run)
        if b and i: t = f"<strong><em>{t}</em></strong>"
        elif b:     t = f"<strong>{t}</strong>"
        elif i:     t = f"<em>{t}</em>"
        parts.append(t)
    if parts:
        return "".join(parts)

    # Fallback: scrape all w:t nodes from raw XML (handles tracked changes, comments, etc.)
    raw = xml_para_text(para)
    if raw:
        return raw.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
    return ""

HEADING_PATTERN = re.compile(
    r"^(Chapter|CHAPTER|Part|PART|פרק|חלק|פרולוג|אפילוג|מבוא|סיום)\b"
)

def is_heading(para, text: str) -> bool:
    style = xml_para_style(para).lower()
    if re.match(r"heading\s*[123]", style): return True
    if re.match(r"כותרת", style):           return True
    if style == "title":                    return True
    if xml_outline_level(para) <= 1:        return True
    if HEADING_PATTERN.match(text):         return True
    return False

def parse_docx(path):
    doc = DocxDocument(path)
    title  = Path(path).stem
    author = ""
    chapters, current_chapter, current_paras = [], None, []
    all_text_sample = []

    print(f"\n[parse_docx] === Starting parse of {Path(path).name} ===")
    print(f"[parse_docx] Total paragraphs in doc: {len(doc.paragraphs)}")

    for i, para in enumerate(doc.paragraphs):
        raw_text   = para.text.strip()           # python-docx native
        xml_text   = xml_para_text(para)          # our XML scraper
        style      = xml_para_style(para)
        outline    = xml_outline_level(para)
        run_count  = len(para.runs)
        run_texts  = [r.text for r in para.runs if r.text]

        # Log every paragraph so we can see what's happening
        print(f"[para {i:03d}] style={style!r:20s} outline={outline} "
              f"runs={run_count} run_texts={run_texts[:3]} "
              f"raw={raw_text[:60]!r} xml={xml_text[:60]!r}")

        text = xml_text or raw_text
        if not text:
            continue

        all_text_sample.append(text)
        style_lower = style.lower()

        if style_lower == "title":
            title = text; continue
        if style_lower == "subtitle":
            author = text; continue

        if is_heading(para, text):
            print(f"  → HEADING detected: {text[:60]!r}")
            if current_chapter is not None:
                chapters.append({"title": current_chapter, "content": current_paras})
                print(f"  → Saved chapter {current_chapter!r} with {len(current_paras)} paras")
            current_chapter = text
            current_paras   = []
        else:
            html = para_to_html(para)
            if html:
                current_paras.append(html)

    if current_chapter is not None:
        chapters.append({"title": current_chapter, "content": current_paras})
        print(f"[parse_docx] Saved final chapter {current_chapter!r} with {len(current_paras)} paras")

    print(f"[parse_docx] Total chapters found: {len(chapters)}")
    for ci, ch in enumerate(chapters):
        print(f"  ch{ci+1}: {ch['title']!r} — {len(ch['content'])} paragraphs")

    # Last-resort fallback
    if not chapters:
        print("[parse_docx] WARNING: no chapters found, using full-doc fallback")
        fallback = [xml_para_text(p) or p.text.strip() for p in doc.paragraphs]
        fallback = [t for t in fallback if t]
        print(f"[parse_docx] Fallback collected {len(fallback)} lines")
        lang_fb = detect_language(" ".join(fallback[:20]))
        label   = "תוכן" if lang_fb == "he" else "Content"
        chapters.append({"title": label, "content": [
            t.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;") for t in fallback
        ]})

    lang = detect_language(" ".join(all_text_sample[:30]))
    print(f"[parse_docx] Detected language: {lang}")
    print(f"[parse_docx] === Done ===\n")
    return {"title": title, "author": author, "chapters": chapters, "detected_lang": lang}


def parse_txt(path):
    title = Path(path).stem
    with open(path, encoding="utf-8", errors="replace") as f:
        lines = f.readlines()

    chapters, current_chapter, current_paras = [], "Content", []
    pat = re.compile(
        r"^(Chapter|CHAPTER|Part|PART|Section|SECTION|פרק|חלק|פרולוג|אפילוג|מבוא)\s*\S*", re.I
    )
    all_text = []

    for line in lines:
        s = line.strip()
        if not s:
            continue
        all_text.append(s)
        is_upper_short = s == s.upper() and 2 < len(s) <= 60 and re.search(r'[A-Z\u05D0-\u05EA]', s)
        if pat.match(s) or is_upper_short:
            if current_paras:
                chapters.append({"title": current_chapter, "content": current_paras})
            current_chapter = s
            current_paras = []
        else:
            current_paras.append(s)

    if current_paras:
        chapters.append({"title": current_chapter, "content": current_paras})

    lang = detect_language(" ".join(all_text[:20]))
    if not chapters:
        chapters.append({"title": "תוכן" if lang == "he" else "Content", "content": all_text})
    return {"title": title, "author": "", "chapters": chapters, "detected_lang": lang}


def parse_md(path):
    title = Path(path).stem
    with open(path, encoding="utf-8", errors="replace") as f:
        lines = f.readlines()

    chapters, current_chapter, current_lines = [], "Content", []
    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
        elif line.startswith("## "):
            if current_lines:
                chapters.append({"title": current_chapter, "content": current_lines, "is_md": True})
            current_chapter = line[3:].strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        chapters.append({"title": current_chapter, "content": current_lines, "is_md": True})
    return {"title": title, "author": "", "chapters": chapters}


# ── EPUB builder ──────────────────────────────────────────────────────────────

def esc(s):
    """HTML-escape a string for safe injection into XHTML."""
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

def build_epub(data, meta, output_path):
    book = epub.EpubBook()
    book.set_identifier(f"id_{uuid.uuid4().hex}")
    book.set_title(meta.get("title") or data["title"])
    book.set_language(meta.get("language") or data.get("detected_lang") or "en")
    if meta.get("author") or data.get("author"):
        book.add_author(meta.get("author") or data["author"])
    if meta.get("publisher"):
        book.add_metadata("DC", "publisher", meta["publisher"])
    if meta.get("description"):
        book.add_metadata("DC", "description", meta["description"])
    if meta.get("subject"):
        for s in [x.strip() for x in meta["subject"].split(",") if x.strip()]:
            book.add_metadata("DC", "subject", s)
    if meta.get("date"):
        book.add_metadata("DC", "date", meta["date"])
    if meta.get("rights"):
        book.add_metadata("DC", "rights", meta["rights"])

    lang = data.get("detected_lang", meta.get("language", "en"))
    is_rtl = lang == "he" or meta.get("language", "") in ("he", "ar", "fa")
    text_align = "right"    if is_rtl else "left"
    css_dir    = "rtl"      if is_rtl else "ltr"
    css_indent = "0"        if is_rtl else "1.5em"
    book.set_direction(css_dir)

    css_content = (
        "body { font-family: 'David Libre', 'Times New Roman', serif;"
        " line-height: 1.85; margin: 5% 8%; color: #222;"
        " direction: " + css_dir + "; text-align: " + text_align + "; }\n"
        "h1 { font-size: 2em; text-align: center; margin: 1.5em 0 0.5em; }\n"
        "h2 { font-size: 1.4em; margin-top: 2em; }\n"
        "p  { margin: 0.5em 0; text-indent: " + css_indent + "; }\n"
        "p.first { text-indent: 0; }\n"
        ".mixed p { unicode-bidi: plaintext; }\n"
    )
    css = epub.EpubItem(uid="style", file_name="style/main.css",
                        media_type="text/css", content=css_content)
    book.add_item(css)

    epub_lang      = meta.get("language") or ("he" if is_rtl else "en")
    body_class     = " class='mixed'" if lang == "mixed" else ""

    epub_chapters = []
    toc = []
    for i, ch in enumerate(data["chapters"]):
        c = epub.EpubHtml(
            title=ch["title"],
            file_name="chapter_%03d.xhtml" % (i+1,),
            lang=epub_lang,
            direction=css_dir,
        )
        ch_title_esc = esc(ch["title"])
        if ch.get("is_md"):
            body = "<h2 class='chapter-title'>" + ch_title_esc + "</h2>\n" + markdown2.markdown("\n".join(ch["content"]))
        else:
            paras = ["<h2 class='chapter-title'>" + ch_title_esc + "</h2>"]
            for j, p in enumerate(ch["content"]):
                paras.append("<p" + (" class='first'" if j == 0 else "") + ">" + p + "</p>")
            body = "\n".join(paras)
        if body_class:
            c.content = "<div" + body_class + ">\n" + body + "\n</div>"
        else:
            c.content = body
        c.add_item(css)
        book.add_item(c)
        epub_chapters.append(c)
        toc.append(epub.Link("chapter_%03d.xhtml" % (i+1,), ch["title"], "ch%d" % (i+1,)))

    book.toc = toc
    book.add_item(epub.EpubNcx())
    nav_title = meta.get("title") or data["title"]
    nav = epub.EpubNav(direction=css_dir, title=nav_title)
    nav.add_item(css)
    book.add_item(nav)
    book.spine = ["nav"] + epub_chapters
    epub.write_epub(output_path, book, {})


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template_string(HTML)

@app.route("/convert", methods=["POST"])
def convert():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Empty filename"}), 400

    ext = Path(f.filename).suffix.lower()
    if ext not in (".docx", ".txt", ".md", ".markdown"):
        return jsonify({"error": f"Unsupported format: {ext}"}), 400

    uid = uuid.uuid4().hex
    src = UPLOAD_FOLDER / f"{uid}{ext}"
    out = UPLOAD_FOLDER / f"{uid}.epub"
    f.save(str(src))

    # Capture all print() output and return it with errors
    import io, contextlib
    log_buf = io.StringIO()
    try:
        with contextlib.redirect_stdout(log_buf):
            if ext == ".docx":
                data = parse_docx(str(src))
            elif ext == ".txt":
                data = parse_txt(str(src))
            else:
                data = parse_md(str(src))

        total_paras = sum(len(c["content"]) for c in data["chapters"])
        if total_paras == 0 and len(data["chapters"]) <= 1:
            src.unlink(missing_ok=True)
            return jsonify({"error": "Document is empty", "log": log_buf.getvalue()}), 400

        meta = {k: request.form.get(k, "").strip() for k in
                ["title","author","language","publisher","date","subject","rights","description"]}
        if not meta["title"]:
            meta["title"] = data["title"]

        with contextlib.redirect_stdout(log_buf):
            build_epub(data, meta, str(out))
        src.unlink(missing_ok=True)

        stem = Path(f.filename).stem
        print(log_buf.getvalue())  # also print to terminal
        return send_file(str(out), as_attachment=True,
                         download_name=f"{stem}.epub",
                         mimetype="application/epub+zip")
    except Exception as e:
        import traceback
        log_buf.write(traceback.format_exc())
        print(log_buf.getvalue())
        src.unlink(missing_ok=True)
        out.unlink(missing_ok=True)
        return jsonify({"error": str(e), "log": log_buf.getvalue()}), 500

@app.route("/preview", methods=["POST"])
def preview():
    """Return detected title, author, chapter list without building epub."""
    if "file" not in request.files:
        return jsonify({"error": "No file"}), 400
    f = request.files["file"]
    ext = Path(f.filename).suffix.lower()
    uid = uuid.uuid4().hex
    src = UPLOAD_FOLDER / f"{uid}{ext}"
    f.save(str(src))
    try:
        if ext == ".docx":   data = parse_docx(str(src))
        elif ext == ".txt":  data = parse_txt(str(src))
        else:                data = parse_md(str(src))
        src.unlink(missing_ok=True)
        print(f"[preview] title={data['title']!r} chapters={len(data['chapters'])} lang={data.get('detected_lang')}")
        for i, ch in enumerate(data['chapters']):
            print(f"  ch{i+1}: {ch['title']!r}  ({len(ch['content'])} paras)")
        return jsonify({
            "title":         data["title"],
            "author":        data.get("author",""),
            "chapters":      [c["title"] for c in data["chapters"]],
            "detected_lang": data.get("detected_lang", "en"),
        })
    except Exception as e:
        src.unlink(missing_ok=True)
        return jsonify({"error": str(e)}), 500


# ── Embedded HTML/CSS/JS UI ───────────────────────────────────────────────────

HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>EPUB Converter</title>
<style>
  *{box-sizing:border-box;margin:0;padding:0}
  body{font-family:'Segoe UI',system-ui,sans-serif;background:#0f1117;color:#e2e8f0;min-height:100vh;display:flex;flex-direction:column;align-items:center;padding:2rem 1rem}
  h1{font-size:1.8rem;font-weight:700;margin-bottom:.25rem;background:linear-gradient(135deg,#818cf8,#c084fc);-webkit-background-clip:text;-webkit-text-fill-color:transparent}
  .subtitle{color:#64748b;font-size:.9rem;margin-bottom:2rem}
  .card{background:#1e2130;border:1px solid #2d3148;border-radius:14px;padding:1.75rem;width:100%;max-width:680px;margin-bottom:1.25rem}
  .card h2{font-size:.85rem;text-transform:uppercase;letter-spacing:.1em;color:#818cf8;margin-bottom:1rem}

  /* Drop zone */
  #dropzone{border:2px dashed #3d4268;border-radius:10px;padding:2.5rem;text-align:center;cursor:pointer;transition:all .2s;background:#151826}
  #dropzone.over{border-color:#818cf8;background:#1a1d35}
  #dropzone.has-file{border-color:#4ade80;background:#0d1f14}
  #dropzone svg{width:40px;height:40px;margin-bottom:.75rem;opacity:.5}
  #dropzone p{color:#64748b;font-size:.9rem}
  #dropzone strong{display:block;color:#e2e8f0;font-size:1rem;margin-bottom:.25rem}
  #file-input{display:none}

  /* Preview box */
  #preview-box{display:none;margin-top:1rem;background:#151826;border-radius:8px;padding:1rem;font-size:.85rem}
  #preview-box .row{display:flex;gap:.5rem;margin-bottom:.3rem}
  #preview-box .lbl{color:#818cf8;min-width:70px}
  #chapter-list{margin-top:.5rem;max-height:120px;overflow-y:auto;padding-left:1rem}
  #chapter-list li{color:#94a3b8;margin:.2rem 0;font-size:.82rem}

  /* Form */
  .grid{display:grid;grid-template-columns:1fr 1fr;gap:.75rem}
  label{display:block;font-size:.78rem;color:#94a3b8;margin-bottom:.3rem}
  input,textarea{width:100%;background:#151826;border:1px solid #2d3148;border-radius:7px;padding:.55rem .75rem;color:#e2e8f0;font-size:.875rem;outline:none;transition:border .15s}
  input:focus,textarea:focus{border-color:#818cf8}
  textarea{resize:vertical;min-height:72px;font-family:inherit}
  .full{grid-column:1/-1}

  /* Button */
  button{display:flex;align-items:center;justify-content:center;gap:.5rem;width:100%;padding:.85rem;border:none;border-radius:10px;font-size:1rem;font-weight:600;cursor:pointer;background:linear-gradient(135deg,#6366f1,#8b5cf6);color:#fff;transition:opacity .2s,transform .1s;margin-top:.25rem}
  button:hover{opacity:.9}
  button:active{transform:scale(.98)}
  button:disabled{opacity:.45;cursor:not-allowed}

  /* Log */
  #log{font-family:'Courier New',monospace;font-size:.8rem;background:#0a0c14;border-radius:8px;padding:1rem;min-height:60px;max-height:160px;overflow-y:auto;color:#4ade80;white-space:pre-wrap;display:none}
  .err{color:#f87171!important}
  .spinner{width:18px;height:18px;border:2px solid rgba(255,255,255,.3);border-top-color:#fff;border-radius:50%;animation:spin .7s linear infinite}
  @keyframes spin{to{transform:rotate(360deg)}}
</style>
</head>
<body>
<h1>EPUB Converter</h1>
<p class="subtitle">Convert DOCX · TXT · Markdown → EPUB</p>

<div class="card">
  <h2>1 — Select File</h2>
  <div id="dropzone" onclick="document.getElementById('file-input').click()">
    <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5">
      <path d="M12 16V4m0 0L8 8m4-4l4 4M4 16v2a2 2 0 002 2h12a2 2 0 002-2v-2"/>
    </svg>
    <strong id="dz-title">Drop file here or click to browse</strong>
    <p id="dz-sub">Supports .docx · .txt · .md · .markdown</p>
  </div>
  <input type="file" id="file-input" accept=".docx,.txt,.md,.markdown"/>
  <div id="preview-box">
    <div class="row"><span class="lbl">Title</span><span id="p-title"></span></div>
    <div class="row"><span class="lbl">Author</span><span id="p-author"></span></div>
    <div class="row"><span class="lbl">Chapters</span><span id="p-count"></span></div>
    <ol id="chapter-list"></ol>
  </div>
</div>

<div class="card">
  <h2>2 — Metadata</h2>
  <div class="grid">
    <div><label>Title</label><input id="m-title" placeholder="Detected from file"/></div>
    <div><label>Author</label><input id="m-author" placeholder="Author name"/></div>
    <div><label>Language</label><input id="m-language" placeholder="en / he / mixed"/></div>
    <div><label>Publisher</label><input id="m-publisher" placeholder="Publisher"/></div>
    <div><label>Date</label><input id="m-date" placeholder="YYYY-MM-DD"/></div>
    <div><label>Subject / Tags</label><input id="m-subject" placeholder="Comma-separated"/></div>
    <div class="full"><label>Rights</label><input id="m-rights" placeholder="© 2024 Author Name"/></div>
    <div class="full"><label>Description</label><textarea id="m-description" placeholder="Short blurb about the book…"></textarea></div>
  </div>
</div>

<div class="card">
  <h2>3 — Convert</h2>
  <div id="log"></div>
  <button id="convert-btn" onclick="doConvert()" disabled>
    <svg viewBox="0 0 24 24" width="18" height="18" fill="none" stroke="currentColor" stroke-width="2">
      <path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5"/>
    </svg>
    Convert & Download EPUB
  </button>
</div>

<script>
let selectedFile = null;

const dropzone = document.getElementById('dropzone');
const fileInput = document.getElementById('file-input');

['dragenter','dragover'].forEach(e => dropzone.addEventListener(e, ev => { ev.preventDefault(); dropzone.classList.add('over'); }));
['dragleave','drop'].forEach(e => dropzone.addEventListener(e, ev => { ev.preventDefault(); dropzone.classList.remove('over'); }));
dropzone.addEventListener('drop', ev => setFile(ev.dataTransfer.files[0]));
fileInput.addEventListener('change', () => setFile(fileInput.files[0]));

function setFile(f) {
  if (!f) return;
  selectedFile = f;
  dropzone.classList.add('has-file');
  document.getElementById('dz-title').textContent = f.name;
  document.getElementById('dz-sub').textContent = (f.size / 1024).toFixed(1) + ' KB';
  document.getElementById('convert-btn').disabled = false;
  // Pre-fill title from filename immediately (preview may override with doc title)
  const stem = f.name.replace(/\.[^.]+$/, '');
  document.getElementById('m-title').value = stem;
  previewFile(f);
}

async function previewFile(f) {
  const fd = new FormData(); fd.append('file', f);
  try {
    const res = await fetch('/preview', {method:'POST', body:fd});
    const d = await res.json();
    if (d.error) return;
    document.getElementById('p-title').textContent  = d.title  || '—';
    document.getElementById('p-author').textContent = d.author || '—';
    document.getElementById('p-count').textContent  = d.chapters.length + ' chapter(s)';
    const ol = document.getElementById('chapter-list');
    ol.innerHTML = d.chapters.map(c => `<li>${c}</li>`).join('');
    document.getElementById('preview-box').style.display = 'block';
    if (!document.getElementById('m-title').value) document.getElementById('m-title').value = d.title;
    if (!document.getElementById('m-author').value && d.author) document.getElementById('m-author').value = d.author;
    if (!document.getElementById('m-language').value && d.detected_lang) document.getElementById('m-language').value = d.detected_lang;
  } catch(e) {}
}

function log(msg, err=false) {
  const el = document.getElementById('log');
  el.style.display = 'block';
  el.innerHTML += `<span class="${err?'err':''}">${msg}</span>\n`;
  el.scrollTop = el.scrollHeight;
}

async function doConvert() {
  if (!selectedFile) return;
  const btn = document.getElementById('convert-btn');
  btn.disabled = true;
  btn.innerHTML = '<div class="spinner"></div> Converting…';
  document.getElementById('log').innerHTML = '';

  const fd = new FormData();
  fd.append('file', selectedFile);
  ['title','author','language','publisher','date','subject','rights','description'].forEach(k => {
    fd.append(k, document.getElementById('m-'+k).value);
  });

  log('Uploading ' + selectedFile.name + ' …');
  try {
    const res = await fetch('/convert', {method:'POST', body:fd});
    if (!res.ok) {
      const d = await res.json();
      log('Error: ' + d.error, true);
      if (d.log) {
        log('-- Server log --');
        d.log.split('\\n').forEach(function(l){ if(l.trim()) log(l); });
      }
    } else {
      log('Parsing & building EPUB …');
      const blob = await res.blob();
      const url  = URL.createObjectURL(blob);
      const a    = document.createElement('a');
      const stem = selectedFile.name.replace(/\.[^.]+$/, '');
      a.href = url; a.download = stem + '.epub'; a.click();
      URL.revokeObjectURL(url);
      log('Done! Download started ✓');
    }
  } catch(e) {
    log('Network error: ' + e, true);
  }

  btn.disabled = false;
  btn.innerHTML = `<svg viewBox="0 0 24 24" width="18" height="18" fill="none" stroke="currentColor" stroke-width="2"><path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5"/></svg> Convert & Download EPUB`;
}
</script>
</body>
</html>
"""

if __name__ == "__main__":
    print("\n  EPUB Converter running at  →  http://localhost:5000\n")
    app.run(debug=False, port=5000)
